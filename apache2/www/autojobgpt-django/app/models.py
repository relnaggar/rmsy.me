from django.db import models
from django.core import files
from django.utils import timezone

import docx
import json
import io
import re
import os

from .scraping import scrape_text
from .gpt import Chat


class IDocumentModel:
  @property
  def has_docx(self):
    return self.docx.name != ""

  @property
  def has_png(self):
    return self.png.name != ""
  
  def generate_png(self):
    if self.has_png:
      raise Exception("the template already has a png file")
    
    # generate the pdf using libreoffice
    outdir = "/".join(self.docx.path.split("/")[:-1])
    os.system(f"libreoffice --headless --convert-to pdf {self.docx.path} " +
      f"--outdir {outdir}")
    pdf_path = self.docx.path.replace(".docx", ".pdf")

    # convert the pdf to png using pdftoppm
    os.system(f"pdftoppm -f 1 -l 1 -png {pdf_path} " + 
      f"{self.docx.path.replace('.docx', '')}")
    png_path = self.docx.path.replace(".docx", "-1.png")

    # remove the pdf file
    os.system(f"rm {pdf_path}")

    # read in the png as a file stream
    png_file_stream = io.BytesIO()
    with open(png_path, "rb") as f:
      png_file_stream.write(f.read())
    
    # remove the extra png file
    os.system(f"rm {png_path}")

    # save the png file stream to the model with the same name as the docx file
    self.png.save(
      png_path.split("/")[-1].replace("-1.png", ".png"),
      files.File(png_file_stream)
    )
  
  def open_document(self):
    # open the template with docx
    return docx.Document(self.docx.path)
  
  def extract_text(self):
    # open the template with docx
    document = self.open_document()
    
    # extract all text from the template
    paragraphs = []
    for paragraph in document.paragraphs:
      paragraphs.append(paragraph.text)
    return "\n".join(paragraphs)


class ResumeTemplateManager(models.Manager):
  def create(self, validated_data):
    resume_template = ResumeTemplate(
      name=validated_data["name"],
      docx=validated_data["docx"],
      description=validated_data["description"]
    )
    resume_template.save()

    fillfield_keys = resume_template.extract_fillfields()
    for key in fillfield_keys:
      if key == "JOB_TITLE":
        description = "The job title."
      elif key == "COMPANY":
        description = "The company name."
      else:
        description = key
      FillField.objects.create(
        key=key,
        description=description,
        template=resume_template,
      )
    
    resume_template.generate_png()
    resume_template.save()
    return resume_template

class ResumeTemplate(models.Model, IDocumentModel):
  objects = ResumeTemplateManager()
  docx = models.FileField(upload_to='templates/')
  png = models.FileField(upload_to='templates/')
  name = models.CharField(max_length=200)
  description = models.TextField(blank=True)

  class Meta:
    constraints = [
      models.UniqueConstraint(fields=['name'], name='unique_name'),
    ]

  def __str__(self):
    return self.name
  
  def extract_fillfields(self, text=None):
    if text is None:
      text = self.extract_text()

    # extract all the fillfields from the template text
    # fillfields must be in the format {{key}}
    return re.findall(r"{{(.*?)}}", text)


class FillField(models.Model):
  template = models.ForeignKey(
    to="ResumeTemplate",
    on_delete=models.CASCADE,
    related_name="fillfields",
  )
  key = models.CharField(max_length=200)
  description = models.TextField(blank=True)

  class Meta:
    constraints = [
      models.UniqueConstraint(fields=['template', 'key'], name='unique_template_key'),
    ]

  def __str__(self):
    return f"{self.key}"


class JobManager(models.Manager):
  def create(self, validated_data):
    job = Job(url=validated_data["url"])
    job.scrape()
    job.extract_job_details()
    job.save()
    return job

class Job(models.Model):
  objects = JobManager()
  url = models.URLField(unique=True)
  title = models.CharField(max_length=26, blank=True)
  company = models.CharField(max_length=26, blank=True)
  text = models.TextField(blank=True)
  chat_messages = models.JSONField(blank=True, null=True)
  date_applied = models.DateTimeField(blank=True, null=True)
  chosen_resume = models.ForeignKey(
    to="Resume", on_delete=models.SET_NULL, null=True, blank=True,
    related_name="chosen_jobs"
  )
  # "backlog", "applying", "pending", "testing"
  # "interviewing", "rejected", "accepted"
  status = models.CharField(max_length=26, blank=True) 

  class Meta:
    constraints = [
      models.UniqueConstraint(
        fields=['title', 'company'],
        name='unique_title_company'),
    ]

  def __str__(self):
    if self.title is None or self.company is None:
      return self.url
    else:
      return f"{self.title}, {self.company}"

  def scrape(self):    
    self.text = scrape_text(self.url)

  def extract_job_details(self):
    chat = Chat()
    response = json.loads(chat.ask(prompt_name="extract_job_details",
      substitutions={"job_text": self.text})
    )
    self.title = response["job_title"]
    self.company = response["company"]
    self.chat_messages = chat.get_additional_messages()

  def apply(self, chosen_resume):
    self.date_applied = timezone.now()
    self.chosen_resume = chosen_resume
    self.status = "pending"
    self.save()


class ResumeManager(models.Manager):
  def create(self, validated_data):
    # if the job already has a resume, then we need to increment the version
    if Resume.objects.filter(job=validated_data["job"]).exists():
      validated_data["version"] = Resume.get_next_version(validated_data["job"])

    resume = Resume(**validated_data)
    resume.save()
    resume.fill()
    resume.generate_docx()
    resume.generate_png()
    resume.save()
    return resume

class Resume(models.Model, IDocumentModel):
  objects = ResumeManager()
  job = models.ForeignKey(to="Job", on_delete=models.SET_NULL, related_name="resumes", null=True)  
  template = models.ForeignKey(to="ResumeTemplate", on_delete=models.SET_NULL, null=True)
  version = models.IntegerField(default=1)
  docx = models.FileField(upload_to='resumes/')
  png = models.FileField(upload_to='resumes/')
  chat_messages = models.JSONField(blank=True, null=True)

  class Meta:
    constraints = [
      models.UniqueConstraint(
        fields=['job', 'version'], name='unique_job_version'
      )
    ]

  def __str__(self):
    return f"{self.job.title}, {self.job.company}, v{self.version}"
  
  @property
  def default_substitutions(self):
    return {
      "JOB_TITLE": self.job.title,
      "COMPANY": self.job.company,
    }

  @property
  def is_filled(self):
    return self.substitutions.count() > 0
  
  @staticmethod
  def get_next_version(job):
    return Resume.objects.filter(job=job).aggregate(
      models.Max('version')
    )["version__max"] + 1

  def _remove_default_substitutions(self, fillfield_keys):
    for default_key in self.default_substitutions.keys():
      if default_key in fillfield_keys:
        fillfield_keys.remove(default_key)

  def _validate_response_and_get_substitutions(self, response, fillfield_keys):
    # validate the response by checking that all the fillfields are present
    substitutions = {}
    for key in fillfield_keys:
      substitutions[key] = response[key]
    return substitutions
  
  def _create_resume_substitutions(self, resume, substitutions):
    # create the resume substitutions
    for key, value in substitutions.items():
      ResumeSubstitution.objects.create(
        resume=resume,
        key=key,
        value=value,
      )

  def _create_new_resume_and_substitutions(self, substitutions, chat_messages):
    # create the new resume
    new_resume = Resume.objects.create(
      job=self.job,
      version=Resume.get_next_version(self.job),
      chat_messages=chat_messages,
    )
    self._create_resume_substitutions(new_resume, substitutions)
    return new_resume

  def fill(self):
    if self.is_filled:
      raise Exception("you can only fill a resume once")
    
    template_text = self.template.extract_text()
    fillfield_keys = self.template.extract_fillfields(text=template_text)
    self._remove_default_substitutions(fillfield_keys)
    
    # construct the fillfields_text part of the prompt    
    fillfields_text = ""
    for key in fillfield_keys:
      fillfield = FillField.objects.get(key=key, template=self.template)
      if fillfield is None:
        raise Exception(f"fillfield {key} not found")
      if fillfield.description == "":
        description = key
      else:
        description = fillfield.description

      fillfields_text +=\
f"""<fillfield>
  <key>{key}</key>
  <description>
{description}
  </description>
</fillfield>\n"""

    # ask GPT to fill in the fillfields
    chat = Chat(self.job.chat_messages)
    response = json.loads(
      chat.ask(prompt_name="fill_resume_template", substitutions={
        "resume_template_text": template_text,
        "fillfields_text": fillfields_text,
      })
    )

    substitutions = self._validate_response_and_get_substitutions(
      response, fillfield_keys
    )

    # save the chat messages 
    self.chat_messages = chat.get_additional_messages()
    self.save()

    self._create_resume_substitutions(self, substitutions)

  def generate_docx(self):
    if not self.is_filled:
      raise Exception(
        "you have to fill the resume with .fill() before you can generate " +
        "the docx file"
      )

    # open the template with docx
    template_document = self.template.open_document()

    # get all the fillfields
    all_substitutions = self.default_substitutions
    for resume_substitution in self.substitutions.all():
      all_substitutions[resume_substitution.key] = resume_substitution.value

    # substitute the fillfields into the template document
    for para in template_document.paragraphs:
      for run in para.runs:
        for key, value in all_substitutions.items():
          if "{{" + key + "}}" in run.text:
            run.text = run.text.replace("{{" + key + "}}", value)

    # save the document to file stream
    file_stream = io.BytesIO()
    template_document.save(file_stream)
    
    # save the file stream to the file field
    file_name = f"{self.job.title}_{self.job.company}_v{self.version}.docx"    
    self.docx.save(file_name, files.File(file_stream))    

  def regenerate(self, feedback=None):
    if not self.is_filled:
      raise Exception(
        "you have to fill the resume with .fill() before you can regenerate " +
        "the resume"
      )
    
    if feedback is None:
      return self._regenerate_without_feedback()
    else:
      return self._regenerate_with_feedback(feedback)

  def _regenerate_without_feedback(self):    
    # the expected fillfields are the keys of the most recent response
    fillfield_keys = list(json.loads(self.chat_messages[-1]['content']).keys())
    self._remove_default_substitutions(fillfield_keys)    

    # re-ask GPT the most recent question
    chat_messages_rewinded_1 = self.chat_messages[:-1]
    chat_message = Chat().ask_with_messages(
      self.job.chat_messages + chat_messages_rewinded_1
    )
    response = json.loads(chat_message.content)

    substitutions = self._validate_response_and_get_substitutions(
      response, fillfield_keys
    )
    new_resume = self._create_new_resume_and_substitutions(
      substitutions, chat_messages_rewinded_1 + [chat_message]
    )

    # copy the resume substitutions from the previous resume
    for key in self.template.extract_fillfields():
      if key not in fillfield_keys and \
          key not in self.default_substitutions.keys():
        # copy the resume substitution
        old_resume_substitution = self.substitutions.get(key=key)
        ResumeSubstitution.objects.create(
          resume=new_resume,
          key=old_resume_substitution.key,
          value=old_resume_substitution.value,
        )
    
    return new_resume

  def _regenerate_with_feedback(self, feedback):
    chat = Chat(self.job.chat_messages + self.chat_messages)
    response = json.loads(
      chat.ask(prompt_name="regenerate_resume", substitutions={
        "feedback": feedback,
      })
    )

    fillfield_keys = self.template.extract_fillfields()
    self._remove_default_substitutions(fillfield_keys)
    substitutions = self._validate_response_and_get_substitutions(
      response, fillfield_keys
    )

    return self._create_new_resume_and_substitutions(
      substitutions, self.chat_messages+chat.get_additional_messages()
    )


class ResumeSubstitution(models.Model):
  resume = models.ForeignKey(to="Resume", on_delete=models.CASCADE, related_name="substitutions")
  key = models.CharField(max_length=200)
  value = models.TextField()

  class Meta:
    constraints = [
      models.UniqueConstraint(fields=['resume', 'key', 'value'], name='unique_resume_key_value'),
    ]

  def __str__(self):
    return f"{self.key}: {self.value} for {self.resume}"
  
  def _regenerate_without_feedback(self):
    # re-ask GPT the most recent question
    chat_messages_rewinded_1 = self.chat_messages[:-1]
    chat_message = Chat().ask_with_messages(
      self.resume.job.chat_messages + self.resume.chat_messages + chat_messages_rewinded_1
    )
    response = json.loads(chat_message.content)
    new_resume = Resume.objects.create(
      job=self.resume.job,
      version=Resume.get_next_version(self.resume.job),
      chat_messages=self.resume.chat_messages + chat_messages_rewinded_1 + [chat_message],
    )
    for resume_substitution in self.resume.substitutions.all():
      if resume_substitution.key != self.key:
        ResumeSubstitution.objects.create(
          resume=new_resume,
          key=resume_substitution.key,
          value=resume_substitution.value,
        )
    return ResumeSubstitution.objects.create(
      resume=new_resume,
      key=self.key,
      value=response[self.key],
    )


  def regenerate(self, feedback):
    chat = Chat(self.resume.job.chat_messages + self.resume.chat_messages)
    response = json.loads(chat.ask(prompt_name="regenerate_substitution", substitutions={
      "key": self.key,
      "feedback": feedback,
    }))

    # copy the resume and its substitutions
    new_resume = Resume.objects.create(
      job=self.resume.job,
      version=Resume.get_next_version(self.resume.job),
      chat_messages=self.resume.chat_messages + chat.get_additional_messages(),
    )
    for resume_substitution in self.resume.substitutions.all():
      if resume_substitution.key != self.key:
        ResumeSubstitution.objects.create(
          resume=new_resume,
          key=resume_substitution.key,
          value=resume_substitution.value,
        )
    return ResumeSubstitution.objects.create(
      resume=new_resume,
      key=self.key,
      value=response[self.key],
    )